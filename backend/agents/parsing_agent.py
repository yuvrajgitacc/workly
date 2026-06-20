import pdfplumber, fitz, os, json, uuid
import threading
from docx import Document
from agents.llm import RotateLLMClient
from pathlib import Path
import re

pymupdf_lock = threading.Lock()

class ResumeParsingAgent:
    def __init__(self):
        self.client = RotateLLMClient()
        self.upload_dir = os.getenv("UPLOAD_DIR", "/tmp/vishleshan/resumes")
        self.photo_dir = os.getenv("PHOTO_DIR", "/tmp/vishleshan/photos")
        os.makedirs(self.photo_dir, exist_ok=True)

    async def parse(self, file_path: str, file_type: str) -> dict:
        try:
            if file_type == "pdf":
                text = self._extract_pdf_text(file_path)
                photo_path = self._extract_pdf_photo(file_path)
            elif file_type in ["docx", "doc"]:
                text = self._extract_docx_text(file_path)
                photo_path = self._extract_docx_photo(file_path)
            else:  # txt
                with open(file_path, 'r', errors='ignore') as f:
                    text = f.read()
                photo_path = None

            if not text or len(text.strip()) < 50:
                return self._empty_result(file_path)

            SCHEMA = """
            {
              "name": string,
              "email": string|null,
              "phone": string|null,
              "location": string|null,
              "linkedin_url": string|null,
              "github_url": string|null,
              "professional_summary": string|null,
              "total_experience_years": float,
              "skills": [
                {"skill":string,"years":float|null,
                 "level":"beginner"|"intermediate"|"expert"|null,
                 "context":string|null}
              ],
              "experience": [
                {"company":string,"role":string,
                 "location":string|null,
                 "start_date":string,"end_date":string,
                 "duration_months":int,
                 "responsibilities":[string],
                 "tech_used":[string]}
              ],
              "education": [
                {"institution":string,"degree":string,
                 "field":string,"year_end":int|null,
                 "cgpa":float|null}
              ],
              "certifications": [
                {"name":string,"issuer":string|null,
                 "year":int|null}
              ],
              "projects": [
                {"name":string,"description":string,
                 "tech_stack":[string],"url":string|null}
              ],
              "achievements": [string],
              "languages": [string]
            }"""

            system = "Expert resume parser. Return ONLY valid JSON. No markdown. No explanation. Null for missing fields."
            prompt = f"Parse resume. Return JSON matching schema:\n{SCHEMA}\n\nResume:\n{text[:8000]}"

            response = self.client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": system},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.1
            )
            
            raw_content = response.choices[0].message.content.strip()
            # Safety strip in case of markdown
            if raw_content.startswith("```json"):
                raw_content = raw_content[7:]
            if raw_content.startswith("```"):
                raw_content = raw_content[3:]
            if raw_content.endswith("```"):
                raw_content = raw_content[:-3]
                
            parsed_dict = json.loads(raw_content.strip())

            email_re = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
            phone_re = r'[\+\(]?[0-9][0-9\s\-\(\)]{8,15}[0-9]'

            if not parsed_dict.get("email"):
                found_emails = re.findall(email_re, text)
                if found_emails:
                    parsed_dict["email"] = found_emails[0]
                    
            if not parsed_dict.get("phone"):
                found_phones = re.findall(phone_re, text)
                if found_phones:
                    parsed_dict["phone"] = found_phones[0].strip()

            return {
                "parsed": parsed_dict,
                "photo_path": photo_path,
                "raw_text_length": len(text),
                "parsing_method": "llm",
                "confidence": 0.9
            }
            
        except Exception as e:
            text = text if 'text' in locals() else ""
            email_re = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
            phone_re = r'[\+\(]?[0-9][0-9\s\-\(\)]{8,15}[0-9]'
            
            email = None
            found_emails = re.findall(email_re, text)
            if found_emails: email = found_emails[0]
            
            phone = None
            found_phones = re.findall(phone_re, text)
            if found_phones: phone = found_phones[0].strip()
            
            fallback_dict = {
                "name": Path(file_path).stem,
                "email": email,
                "phone": phone,
                "location": None,
                "skills": [],
                "experience": [],
                "education": [],
                "total_experience_years": 0
            }
            
            return {
                "parsed": fallback_dict,
                "photo_path": photo_path if 'photo_path' in locals() else None,
                "raw_text_length": len(text),
                "parsing_method": "fallback",
                "confidence": 0.4
            }

    def _extract_pdf_text(self, path):
        with pdfplumber.open(path) as pdf:
            pages = [p.extract_text() or "" for p in pdf.pages]
        return "\n".join(pages)

    def _extract_pdf_photo(self, path):
        try:
            with pymupdf_lock:
                doc = fitz.open(path)
                for page in doc:
                    for img in page.get_images():
                        xref = img[0]
                        base = doc.extract_image(xref)
                        img_bytes = base["image"]
                        photo_path = f"{self.photo_dir}/{uuid.uuid4()}.jpg"
                        with open(photo_path, "wb") as f: 
                            f.write(img_bytes)
                        return photo_path
        except: 
            return None

    def _extract_docx_text(self, path):
        doc = Document(path)
        parts = []
        for para in doc.paragraphs:
            if para.text.strip(): 
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    c.text.strip() for c in row.cells if c.text.strip()
                )
                if row_text: 
                    parts.append(row_text)
        return "\n".join(parts)

    def _extract_docx_photo(self, path):
        try:
            from docx.oxml.ns import qn
            doc = Document(path)
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    img_data = rel.target_part.blob
                    photo_path = f"{self.photo_dir}/{uuid.uuid4()}.jpg"
                    with open(photo_path, "wb") as f: 
                        f.write(img_data)
                    return photo_path
        except: 
            return None

    def _empty_result(self, path):
        return {
            "parsed": {
                "name": Path(path).stem,
                "email": None,
                "phone": None,
                "location": None,
                "skills": [],
                "experience": [],
                "education": [],
                "total_experience_years": 0
            },
            "photo_path": None,
            "raw_text_length": 0,
            "parsing_method": "empty",
            "confidence": 0.0
        }
